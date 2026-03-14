"""
Report Generation Module.

Generates PDF reports using ReportLab for:
- Combined review reports (all Stage 1 + Stage 2 reviews for a single proposal)
- Cycle summary reports (status counts, acceptance rates, reviewer workload)

Both functions return a BytesIO buffer so the caller can stream it as an
HTTP response without writing temp files to disk.
"""
import io
import logging
from xml.sax.saxutils import escape

from django.core.exceptions import ObjectDoesNotExist
from django.db.models import Count, Q
from django.utils import timezone


logger = logging.getLogger(__name__)


def _safe_text(value):
    """Escape dynamic values before inserting into ReportLab Paragraph markup.

    ReportLab's Paragraph uses XML-like markup, so unescaped user input
    (e.g., angle brackets in a title) would break the PDF rendering.
    """
    if value is None:
        return ""
    return escape(str(value))


def _canonical_cycle_status_rows(proposals):
    """Build status rows using the mandatory lifecycle status list."""
    labels = {
        'SUBMITTED': 'Submitted',
        'UNDER_STAGE_1_REVIEW': 'Under Stage 1 Review',
        'STAGE_1_REJECTED': 'Stage 1 Rejected',
        'ACCEPTED_NO_CORRECTIONS': 'Accepted (No Corrections)',
        'TENTATIVELY_ACCEPTED': 'Tentatively Accepted',
        'REVISION_REQUESTED': 'Revision Requested',
        'REVISED_PROPOSAL_SUBMITTED': 'Revised Proposal Submitted',
        'UNDER_STAGE_2_REVIEW': 'Under Stage 2 Review',
        'FINAL_ACCEPTED': 'Final Accepted',
        'FINAL_REJECTED': 'Final Rejected',
    }

    counts = {status: 0 for status in proposals.model.CANONICAL_LIFECYCLE_STATUSES}
    for row in proposals.values('status').annotate(total=Count('id')):
        reportable = proposals.model.reportable_status(row['status'])
        if reportable:
            counts[reportable] += row['total']

    rows = [('Total Proposals', sum(counts.values()))]
    for status in proposals.model.CANONICAL_LIFECYCLE_STATUSES:
        rows.append((labels[status], counts[status]))
    return rows, counts


def generate_combined_review_pdf(proposal):
    """
    Generate a combined PDF report with all reviews for a proposal.
    
    Args:
        proposal: Proposal instance
        
    Returns:
        BytesIO buffer containing the PDF
    """
    # Lazy import — reportlab is a large library, so only load it when actually
    # generating a PDF. This also provides a graceful degradation message if
    # the dependency is missing in a lightweight deployment.
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except ImportError:
        buffer = io.BytesIO()
        buffer.write(b"Report generation requires reportlab. Please install it with: pip install reportlab")
        buffer.seek(0)
        return buffer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)

    # ReportLab uses a "story" list — each element is rendered sequentially in the PDF
    story = []
    styles = getSampleStyleSheet()

    # --- Report Title ---
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    story.append(Paragraph("Combined Review Report", title_style))
    story.append(Spacer(1, 12))

    # --- Proposal Metadata Section ---
    fund_requested = proposal.fund_requested
    if fund_requested is None:
        fund_requested_text = "N/A"
    else:
        fund_requested_text = f"${fund_requested:,.2f}"
    story.append(Paragraph(f"<b>Proposal Code:</b> {_safe_text(proposal.proposal_code)}", styles['Normal']))
    story.append(Paragraph(f"<b>Title:</b> {_safe_text(proposal.title)}", styles['Normal']))
    story.append(Paragraph(f"<b>PI:</b> {_safe_text(proposal.pi_name)}", styles['Normal']))
    story.append(Paragraph(f"<b>Department:</b> {_safe_text(proposal.pi_department)}", styles['Normal']))
    if proposal.co_investigators:
        story.append(Paragraph(f"<b>Co-Investigators:</b> {_safe_text(proposal.co_investigators)}", styles['Normal']))
    story.append(Paragraph(f"<b>Funding Requested:</b> {_safe_text(fund_requested_text)}", styles['Normal']))
    story.append(Paragraph(f"<b>Status:</b> {_safe_text(proposal.get_status_display())}", styles['Normal']))

    if proposal.abstract:
        story.append(Spacer(1, 10))
        story.append(Paragraph("<b>Abstract:</b>", styles['Normal']))
        story.append(Paragraph(_safe_text(proposal.abstract), styles['Normal']))

    story.append(Spacer(1, 20))
    
    # --- Stage 1 Reviews Section ---
    # Each completed Stage 1 assignment gets a score table showing all 8 criteria
    # with their individual scores, plus the computed total and percentage.
    story.append(Paragraph("Stage 1 Reviews", styles['Heading2']))
    story.append(Spacer(1, 10))

    from reviews.models import ReviewAssignment, Stage2Review
    stage1_assignments = ReviewAssignment.objects.filter(
        proposal=proposal,
        stage=ReviewAssignment.Stage.STAGE_1,
        status=ReviewAssignment.Status.COMPLETED
    ).select_related('stage1_score')

    for i, assignment in enumerate(stage1_assignments, 1):
        # Reviewers are numbered anonymously (Reviewer 1, 2, ...) to preserve blinding
        story.append(Paragraph(f"<b>Reviewer {i}</b>", styles['Heading3']))

        try:
            score = assignment.stage1_score

            # 8-criteria scoring table matching the Stage1Score model fields.
            # Max scores: 5 criteria × 15pts + 2 × 10pts + 1 × 5pts = 100 total.
            score_data = [
                ['Criteria', 'Score', 'Max'],
                ['Originality', str(score.originality_score), '15'],
                ['Clarity', str(score.clarity_score), '15'],
                ['Literature Review', str(score.literature_review_score), '15'],
                ['Methodology', str(score.methodology_score), '15'],
                ['Impact', str(score.impact_score), '15'],
                ['Publication Potential', str(score.publication_potential_score), '10'],
                ['Budget Appropriateness', str(score.budget_appropriateness_score), '10'],
                ['Timeline Practicality', str(score.timeline_practicality_score), '5'],
                ['Total', str(score.total_score), '100'],
                ['Percentage', f"{score.percentage_score}%", '-'],
                ['Weighted Score', f"{score.weighted_percentage_score}%", '-'],
            ]

            table = Table(score_data, colWidths=[2.5*inch, 1*inch, 0.75*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),        # Header row styling
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, -2), (-1, -1), colors.beige),     # Total/percentage rows highlighted
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
            story.append(Spacer(1, 10))
            
            if score.narrative_comments:
                story.append(Paragraph(f"<b>Comments:</b>", styles['Normal']))
                story.append(Paragraph(_safe_text(score.narrative_comments), styles['Normal']))
            if score.recommendation:
                story.append(Paragraph(
                    f"<b>Recommendation:</b> {_safe_text(score.get_recommendation_display())}",
                    styles['Normal']
                ))
            if score.detailed_recommendation:
                story.append(Paragraph("<b>Detailed Recommendation:</b>", styles['Normal']))
                story.append(Paragraph(_safe_text(score.detailed_recommendation), styles['Normal']))
            
        except ObjectDoesNotExist:
            story.append(Paragraph("Score data not available", styles['Normal']))
        except Exception:
            logger.exception("Unexpected error while rendering stage 1 score for assignment_id=%s", assignment.id)
            story.append(Paragraph("Score data not available", styles['Normal']))
        
        story.append(Spacer(1, 15))
    
    if not stage1_assignments.exists():
        story.append(Paragraph("No Stage 1 reviews completed yet.", styles['Normal']))
    
    # --- Stage 2 Reviews Section ---
    # Stage 2 only applies to proposals that were tentatively accepted or had
    # revisions requested. These reviews assess whether the PI addressed concerns.
    story.append(Spacer(1, 20))
    story.append(Paragraph("Stage 2 Reviews", styles['Heading2']))
    story.append(Spacer(1, 10))

    stage2_assignments = ReviewAssignment.objects.filter(
        proposal=proposal,
        stage=ReviewAssignment.Stage.STAGE_2,
        status=ReviewAssignment.Status.COMPLETED
    ).select_related('stage2_review')
    
    for i, assignment in enumerate(stage2_assignments, 1):
        story.append(Paragraph(f"<b>Reviewer {i}</b>", styles['Heading3']))
        
        try:
            review = assignment.stage2_review
            story.append(Paragraph(f"<b>Concerns Addressed:</b> {_safe_text(review.get_concerns_addressed_display())}", styles['Normal']))
            story.append(Paragraph(f"<b>Recommendation:</b> {_safe_text(review.get_revised_recommendation_display())}", styles['Normal']))
            if review.revised_score is not None:
                story.append(Paragraph(f"<b>Revised Score:</b> {_safe_text(review.revised_score)}%", styles['Normal']))
            if review.technical_comments:
                story.append(Paragraph(f"<b>Technical Comments:</b>", styles['Normal']))
                story.append(Paragraph(_safe_text(review.technical_comments), styles['Normal']))
            if review.budget_comments:
                story.append(Paragraph(f"<b>Budget Comments:</b>", styles['Normal']))
                story.append(Paragraph(_safe_text(review.budget_comments), styles['Normal']))
        except ObjectDoesNotExist:
            story.append(Paragraph("Review data not available", styles['Normal']))
        except Exception:
            logger.exception("Unexpected error while rendering stage 2 review for assignment_id=%s", assignment.id)
            story.append(Paragraph("Review data not available", styles['Normal']))
        
        story.append(Spacer(1, 15))
    
    chair_stage2_reviews = Stage2Review.objects.filter(
        proposal=proposal,
        is_chair_review=True,
        is_draft=False
    ).select_related('reviewed_by')

    for review in chair_stage2_reviews:
        author_name = review.reviewed_by.get_full_name() if review.reviewed_by else 'SRC Chair'
        story.append(Paragraph(f"<b>{_safe_text(author_name or 'SRC Chair')}</b>", styles['Heading3']))
        story.append(Paragraph(f"<b>Concerns Addressed:</b> {_safe_text(review.get_concerns_addressed_display())}", styles['Normal']))
        story.append(Paragraph(f"<b>Recommendation:</b> {_safe_text(review.get_revised_recommendation_display())}", styles['Normal']))
        if review.revised_score is not None:
            story.append(Paragraph(f"<b>Revised Score:</b> {_safe_text(review.revised_score)}%", styles['Normal']))
        if review.technical_comments:
            story.append(Paragraph(f"<b>Technical Comments:</b>", styles['Normal']))
            story.append(Paragraph(_safe_text(review.technical_comments), styles['Normal']))
        if review.budget_comments:
            story.append(Paragraph(f"<b>Budget Comments:</b>", styles['Normal']))
            story.append(Paragraph(_safe_text(review.budget_comments), styles['Normal']))
        story.append(Spacer(1, 15))

    if not stage2_assignments.exists() and not chair_stage2_reviews.exists():
        story.append(Paragraph("No Stage 2 reviews completed yet.", styles['Normal']))

    # --- Decisions Section ---
    from proposals.models import Stage1Decision, FinalDecision

    try:
        s1_decision = proposal.stage1_decision
        story.append(Spacer(1, 20))
        story.append(Paragraph("Stage 1 Decision", styles['Heading2']))
        story.append(Spacer(1, 8))
        story.append(Paragraph(f"<b>Decision:</b> {_safe_text(s1_decision.get_decision_display())}", styles['Normal']))
        if s1_decision.chair_comments:
            story.append(Paragraph(f"<b>Chair Comments:</b> {_safe_text(s1_decision.chair_comments)}", styles['Normal']))
        story.append(Paragraph(f"<b>Date:</b> {s1_decision.decision_date.strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    except (ObjectDoesNotExist, Stage1Decision.DoesNotExist):
        pass

    try:
        final = proposal.final_decision
        story.append(Spacer(1, 20))
        story.append(Paragraph("Final Decision", styles['Heading2']))
        story.append(Spacer(1, 8))
        story.append(Paragraph(f"<b>Decision:</b> {_safe_text(final.get_decision_display())}", styles['Normal']))
        if final.approved_grant_amount is not None:
            story.append(Paragraph(f"<b>Approved Amount:</b> ${final.approved_grant_amount:,.2f}", styles['Normal']))
        if final.final_remarks:
            story.append(Paragraph(f"<b>Remarks:</b> {_safe_text(final.final_remarks)}", styles['Normal']))
        story.append(Paragraph(f"<b>Date:</b> {final.decision_date.strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    except (ObjectDoesNotExist, FinalDecision.DoesNotExist):
        pass

    # --- Footer with generation timestamp ---
    story.append(Spacer(1, 30))
    story.append(Paragraph(
        f"Generated on {timezone.localtime().strftime('%Y-%m-%d %H:%M:%S')}",
        ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
    ))

    # Build the PDF into the in-memory buffer and rewind for reading
    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_summary_report(cycle):
    """
    Generate a summary report for a grant cycle.

    Args:
        cycle: GrantCycle instance

    Returns:
        BytesIO buffer containing the PDF
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except ImportError:
        buffer = io.BytesIO()
        buffer.write(b"Report generation requires reportlab.")
        buffer.seek(0)
        return buffer

    from proposals.models import Proposal, Stage1Decision, FinalDecision
    from reviews.models import ReviewAssignment, ReviewerProfile

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()

    story.append(Paragraph("CTRG Grant Cycle Summary Report", styles['Heading1']))
    story.append(Paragraph(f"{_safe_text(cycle.name)} ({_safe_text(cycle.year)})", styles['Heading2']))
    story.append(Spacer(1, 20))

    proposals = cycle.proposals.all()
    status_rows, _ = _canonical_cycle_status_rows(proposals.exclude(status=Proposal.Status.DRAFT))

    stage1_decided = proposals.filter(stage1_decision__isnull=False).count()
    stage1_accepted = proposals.filter(
        stage1_decision__decision__in=[
            Stage1Decision.Decision.ACCEPT,
            Stage1Decision.Decision.TENTATIVELY_ACCEPT,
        ]
    ).count()
    s1_rate = f"{(stage1_accepted / stage1_decided * 100):.1f}%" if stage1_decided > 0 else "N/A"

    final_decided = proposals.filter(final_decision__isnull=False).count()
    final_accepted = proposals.filter(final_decision__decision=FinalDecision.Decision.ACCEPTED).count()
    final_rate = f"{(final_accepted / final_decided * 100):.1f}%" if final_decided > 0 else "N/A"

    stats = [['Status', 'Count']] + [[label, str(value)] for label, value in status_rows]
    table = Table(stats, colWidths=[3 * inch, 1.5 * inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    story.append(table)
    story.append(Spacer(1, 15))

    story.append(Paragraph("Acceptance Rates", styles['Heading2']))
    story.append(Spacer(1, 8))
    rates_data = [
        ['Metric', 'Value'],
        ['Stage 1 Acceptance Rate', s1_rate],
        ['Final Acceptance Rate', final_rate],
    ]
    rates_table = Table(rates_data, colWidths=[3 * inch, 1.5 * inch])
    rates_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
        ('FONTNAME', (1, 1), (-1, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    story.append(rates_table)
    story.append(Spacer(1, 20))

    story.append(Paragraph("Reviewer Workload Summary", styles['Heading2']))
    story.append(Spacer(1, 8))

    reviewer_profiles = ReviewerProfile.objects.select_related('user').filter(
        user__review_assignments__proposal__cycle=cycle
    ).annotate(
        s1_count=Count('user__review_assignments', filter=Q(
            user__review_assignments__stage=ReviewAssignment.Stage.STAGE_1,
            user__review_assignments__proposal__cycle=cycle,
        )),
        s2_count=Count('user__review_assignments', filter=Q(
            user__review_assignments__stage=ReviewAssignment.Stage.STAGE_2,
            user__review_assignments__proposal__cycle=cycle,
        )),
        total=Count('user__review_assignments', filter=Q(
            user__review_assignments__proposal__cycle=cycle,
        )),
        pending=Count('user__review_assignments', filter=Q(
            user__review_assignments__status=ReviewAssignment.Status.PENDING,
            user__review_assignments__proposal__cycle=cycle,
        )),
    ).distinct()

    workload_data = [['Reviewer', 'Department', 'Stage 1', 'Stage 2', 'Total', 'Pending']]
    for profile in reviewer_profiles:
        if profile.total <= 0:
            continue
        reviewer_name = profile.user.get_full_name() or profile.user.username
        workload_data.append([
            _safe_text(reviewer_name),
            _safe_text(profile.department),
            str(profile.s1_count),
            str(profile.s2_count),
            str(profile.total),
            str(profile.pending),
        ])

    if len(workload_data) > 1:
        workload_table = Table(
            workload_data,
            colWidths=[1.8 * inch, 1.2 * inch, 0.6 * inch, 0.6 * inch, 0.6 * inch, 0.7 * inch]
        )
        workload_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (2, 1), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(workload_table)
    else:
        story.append(Paragraph("No reviewer assignments for this cycle.", styles['Normal']))

    story.append(Spacer(1, 20))
    story.append(Paragraph(
        f"Generated on {timezone.localtime().strftime('%Y-%m-%d %H:%M:%S')}",
        ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
    ))

    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_review_template_pdf(proposal):
    """
    Generate a reviewer scoring template PDF for a proposal.

    Produces a printable form with proposal metadata and blank scoring
    rubric so reviewers have a reference document for offline evaluation.

    Args:
        proposal: Proposal instance

    Returns:
        BytesIO buffer containing the PDF
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except ImportError:
        buffer = io.BytesIO()
        buffer.write(b"Report generation requires reportlab.")
        buffer.seek(0)
        return buffer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)

    story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
        alignment=1
    )
    story.append(Paragraph("CTRG Review Scoring Template", title_style))
    story.append(Spacer(1, 12))

    # Proposal metadata
    fund_requested = proposal.fund_requested
    fund_text = f"${fund_requested:,.2f}" if fund_requested is not None else "N/A"

    story.append(Paragraph(f"<b>Proposal Code:</b> {_safe_text(proposal.proposal_code)}", styles['Normal']))
    story.append(Paragraph(f"<b>Title:</b> {_safe_text(proposal.title)}", styles['Normal']))
    story.append(Paragraph(f"<b>PI:</b> {_safe_text(proposal.pi_name)}", styles['Normal']))
    story.append(Paragraph(f"<b>Department:</b> {_safe_text(proposal.pi_department)}", styles['Normal']))
    story.append(Paragraph(f"<b>Funding Requested:</b> {_safe_text(fund_text)}", styles['Normal']))

    if proposal.abstract:
        story.append(Spacer(1, 10))
        story.append(Paragraph("<b>Abstract:</b>", styles['Normal']))
        story.append(Paragraph(_safe_text(proposal.abstract), styles['Normal']))

    if proposal.co_investigators:
        story.append(Paragraph(f"<b>Co-Investigators:</b> {_safe_text(proposal.co_investigators)}", styles['Normal']))

    story.append(Spacer(1, 20))

    # Scoring rubric table with blank score column
    story.append(Paragraph("Stage 1 Scoring Rubric", styles['Heading2']))
    story.append(Spacer(1, 10))

    rubric_data = [
        ['Criteria', 'Max Score', 'Your Score', 'Comments'],
        ['Originality', '15', '', ''],
        ['Clarity of Research Problem', '15', '', ''],
        ['Literature Review', '15', '', ''],
        ['Methodology', '15', '', ''],
        ['Potential Impact', '15', '', ''],
        ['Publication Potential', '10', '', ''],
        ['Budget Appropriateness', '10', '', ''],
        ['Timeline Practicality', '5', '', ''],
        ['Total', '100', '', ''],
    ]

    table = Table(rubric_data, colWidths=[2*inch, 0.8*inch, 0.8*inch, 2.5*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, -1), (-1, -1), colors.beige),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.Color(0.95, 0.95, 0.95)]),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
    ]))
    story.append(table)
    story.append(Spacer(1, 20))

    # Recommendation section
    story.append(Paragraph("Recommendation", styles['Heading2']))
    story.append(Spacer(1, 8))
    rec_data = [
        ['[ ] Accept without corrections'],
        ['[ ] Tentatively accept (revisions needed)'],
        ['[ ] Reject'],
    ]
    rec_table = Table(rec_data, colWidths=[5*inch])
    rec_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(rec_table)
    story.append(Spacer(1, 15))

    # Narrative comments section
    story.append(Paragraph("Narrative Comments:", styles['Heading3']))
    story.append(Spacer(1, 8))
    for _ in range(8):
        story.append(Paragraph("_" * 90, styles['Normal']))
        story.append(Spacer(1, 4))

    story.append(Spacer(1, 20))
    story.append(Paragraph("Reviewer Signature: ________________    Date: ________________", styles['Normal']))

    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph(
        f"Generated on {timezone.localtime().strftime('%Y-%m-%d %H:%M:%S')}",
        ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
    ))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _build_docx_document(title):
    from docx import Document

    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(f"Generated on {timezone.localtime().strftime('%Y-%m-%d %H:%M:%S')}")
    return document


def _docx_buffer(document):
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def generate_combined_review_docx(proposal):
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from reviews.models import ReviewAssignment, Stage2Review

    document = _build_docx_document("Combined Review Report")
    document.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading("Proposal Details", level=2)
    details = [
        ("Proposal Code", proposal.proposal_code),
        ("Title", proposal.title),
        ("PI", proposal.pi_name),
        ("Department", proposal.pi_department),
        ("Funding Requested", f"${proposal.fund_requested:,.2f}" if proposal.fund_requested is not None else "N/A"),
        ("Status", proposal.get_status_display()),
    ]
    for label, value in details:
        document.add_paragraph(f"{label}: {value}")
    if proposal.abstract:
        document.add_paragraph(f"Abstract: {proposal.abstract}")

    document.add_heading("Stage 1 Reviews", level=2)
    stage1_assignments = ReviewAssignment.objects.filter(
        proposal=proposal,
        stage=ReviewAssignment.Stage.STAGE_1,
        status=ReviewAssignment.Status.COMPLETED,
    ).select_related('stage1_score', 'reviewer')
    if not stage1_assignments.exists():
        document.add_paragraph("No Stage 1 reviews completed yet.")
    for index, assignment in enumerate(stage1_assignments, start=1):
        if not hasattr(assignment, 'stage1_score'):
            continue
        score = assignment.stage1_score
        document.add_heading(f"Reviewer {index}", level=3)
        document.add_paragraph(f"Recommendation: {score.get_recommendation_display() if score.recommendation else 'N/A'}")
        rubric_table = document.add_table(rows=1, cols=3)
        rubric_table.rows[0].cells[0].text = "Criteria"
        rubric_table.rows[0].cells[1].text = "Score"
        rubric_table.rows[0].cells[2].text = "Max"
        rows = [
            ("Originality", score.originality_score, 15),
            ("Clarity", score.clarity_score, 15),
            ("Literature Review", score.literature_review_score, 15),
            ("Methodology", score.methodology_score, 15),
            ("Impact", score.impact_score, 15),
            ("Publication Potential", score.publication_potential_score, 10),
            ("Budget Appropriateness", score.budget_appropriateness_score, 10),
            ("Timeline Practicality", score.timeline_practicality_score, 5),
            ("Total", score.total_score, 100),
        ]
        for label, value, max_score in rows:
            row = rubric_table.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value)
            row[2].text = str(max_score)
        document.add_paragraph(f"Narrative Comments: {score.narrative_comments or 'N/A'}")
        if score.detailed_recommendation:
            document.add_paragraph(f"Detailed Recommendation: {score.detailed_recommendation}")

    document.add_heading("Stage 2 Reviews", level=2)
    stage2_reviews = []
    for assignment in ReviewAssignment.objects.filter(
        proposal=proposal,
        stage=ReviewAssignment.Stage.STAGE_2,
        status=ReviewAssignment.Status.COMPLETED,
    ).select_related('stage2_review', 'reviewer'):
        if hasattr(assignment, 'stage2_review'):
            stage2_reviews.append((assignment.reviewer.get_full_name() or assignment.reviewer.username, assignment.stage2_review))
    for chair_review in Stage2Review.objects.filter(proposal=proposal, is_chair_review=True, is_draft=False).select_related('reviewed_by'):
        author = chair_review.reviewed_by.get_full_name() if chair_review.reviewed_by else 'SRC Chair'
        stage2_reviews.append((author or 'SRC Chair', chair_review))

    if not stage2_reviews:
        document.add_paragraph("No Stage 2 reviews completed yet.")
    for author, review in stage2_reviews:
        document.add_heading(author, level=3)
        document.add_paragraph(f"Concerns Addressed: {review.get_concerns_addressed_display()}")
        document.add_paragraph(f"Recommendation: {review.get_revised_recommendation_display()}")
        if review.revised_score is not None:
            document.add_paragraph(f"Revised Score: {review.revised_score}%")
        document.add_paragraph(f"Technical Comments: {review.technical_comments or 'N/A'}")
        if review.budget_comments:
            document.add_paragraph(f"Budget Comments: {review.budget_comments}")

    if hasattr(proposal, 'stage1_decision'):
        document.add_heading("Stage 1 Decision", level=2)
        document.add_paragraph(f"Decision: {proposal.stage1_decision.get_decision_display()}")
        document.add_paragraph(f"Average Score: {proposal.stage1_decision.average_score}")
        if proposal.stage1_decision.chair_comments:
            document.add_paragraph(f"Chair Comments: {proposal.stage1_decision.chair_comments}")
    if hasattr(proposal, 'final_decision'):
        document.add_heading("Final Decision", level=2)
        document.add_paragraph(f"Decision: {proposal.final_decision.get_decision_display()}")
        document.add_paragraph(f"Approved Amount: ${proposal.final_decision.approved_grant_amount:,.2f}")
        document.add_paragraph(f"Final Remarks: {proposal.final_decision.final_remarks}")

    return _docx_buffer(document)


def generate_review_template_docx(proposal):
    document = _build_docx_document("CTRG Review Scoring Template")
    document.add_heading("Proposal Details", level=2)
    for line in (
        f"Proposal Code: {proposal.proposal_code}",
        f"Title: {proposal.title}",
        f"PI: {proposal.pi_name}",
        f"Department: {proposal.pi_department}",
        f"Funding Requested: ${proposal.fund_requested:,.2f}" if proposal.fund_requested is not None else "Funding Requested: N/A",
    ):
        document.add_paragraph(line)
    if proposal.abstract:
        document.add_paragraph(f"Abstract: {proposal.abstract}")

    document.add_heading("Stage 1 Scoring Rubric", level=2)
    table = document.add_table(rows=1, cols=4)
    for cell, header in zip(table.rows[0].cells, ["Criteria", "Max Score", "Your Score", "Comments"]):
        cell.text = header
    rubric_rows = [
        ("Originality", "15"),
        ("Clarity of Research Problem", "15"),
        ("Literature Review", "15"),
        ("Methodology", "15"),
        ("Potential Impact", "15"),
        ("Publication Potential", "10"),
        ("Budget Appropriateness", "10"),
        ("Timeline Practicality", "5"),
        ("Total", "100"),
    ]
    for label, max_score in rubric_rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = max_score
        row[2].text = ""
        row[3].text = ""
    document.add_heading("Recommendation", level=2)
    for option in (
        "[ ] Accept without corrections",
        "[ ] Tentatively accept (revisions needed)",
        "[ ] Reject",
    ):
        document.add_paragraph(option)
    document.add_paragraph("Narrative Comments:")
    for _ in range(6):
        document.add_paragraph("____________________________________________________________")
    return _docx_buffer(document)


def generate_summary_report_docx(cycle):
    from proposals.models import Proposal, Stage1Decision, FinalDecision
    from reviews.models import ReviewAssignment, ReviewerProfile

    document = _build_docx_document("CTRG Grant Cycle Summary Report")
    document.add_paragraph(f"Cycle: {cycle.name} ({cycle.year})")

    proposals = cycle.proposals.all()
    status_rows, _ = _canonical_cycle_status_rows(proposals.exclude(status=Proposal.Status.DRAFT))
    document.add_heading("Status Breakdown", level=2)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Status"
    table.rows[0].cells[1].text = "Count"
    for label, value in status_rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    stage1_decided = proposals.filter(stage1_decision__isnull=False).count()
    stage1_accepted = proposals.filter(
        stage1_decision__decision__in=[
            Stage1Decision.Decision.ACCEPT,
            Stage1Decision.Decision.TENTATIVELY_ACCEPT,
        ]
    ).count()
    final_decided = proposals.filter(final_decision__isnull=False).count()
    final_accepted = proposals.filter(final_decision__decision=FinalDecision.Decision.ACCEPTED).count()
    document.add_heading("Acceptance Rates", level=2)
    document.add_paragraph(
        f"Stage 1 Acceptance Rate: {(stage1_accepted / stage1_decided * 100):.1f}%" if stage1_decided else "Stage 1 Acceptance Rate: N/A"
    )
    document.add_paragraph(
        f"Final Acceptance Rate: {(final_accepted / final_decided * 100):.1f}%" if final_decided else "Final Acceptance Rate: N/A"
    )

    document.add_heading("Reviewer Workloads", level=2)
    workload_table = document.add_table(rows=1, cols=6)
    headers = ["Reviewer", "Department", "Stage 1", "Stage 2", "Total", "Pending"]
    for cell, header in zip(workload_table.rows[0].cells, headers):
        cell.text = header
    reviewer_profiles = ReviewerProfile.objects.select_related('user').filter(
        user__review_assignments__proposal__cycle=cycle
    ).annotate(
        s1_count=Count('user__review_assignments', filter=Q(
            user__review_assignments__stage=ReviewAssignment.Stage.STAGE_1,
            user__review_assignments__proposal__cycle=cycle,
        )),
        s2_count=Count('user__review_assignments', filter=Q(
            user__review_assignments__stage=ReviewAssignment.Stage.STAGE_2,
            user__review_assignments__proposal__cycle=cycle,
        )),
        total=Count('user__review_assignments', filter=Q(user__review_assignments__proposal__cycle=cycle)),
        pending=Count('user__review_assignments', filter=Q(
            user__review_assignments__status=ReviewAssignment.Status.PENDING,
            user__review_assignments__proposal__cycle=cycle,
        )),
    ).distinct()
    for profile in reviewer_profiles:
        if profile.total <= 0:
            continue
        row = workload_table.add_row().cells
        row[0].text = profile.user.get_full_name() or profile.user.username
        row[1].text = profile.department or ''
        row[2].text = str(profile.s1_count)
        row[3].text = str(profile.s2_count)
        row[4].text = str(profile.total)
        row[5].text = str(profile.pending)

    return _docx_buffer(document)
